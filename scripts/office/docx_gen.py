#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档生成器 — 接收结构化 JSON 描述，用 python-docx 渲染 .docx 文件。

用法:
  python docx_gen.py <input.json> <output_dir>

input.json 结构:
{
  "title": "文档标题",
  "author": "作者（可选）",
  "intro": "文档简介（可选，显示在标题下方）",
  "blocks": [
    { "type": "h1",   "text": "一级标题" },
    { "type": "h2",   "text": "二级标题" },
    { "type": "h3",   "text": "三级标题" },
    { "type": "p",    "text": "普通段落" },
    { "type": "bullet", "text": "项目符号" },
    { "type": "number", "text": "编号项" },
    { "type": "quote", "text": "引用" }
  ]
}

输出: 打印 JSON 到 stdout:
  {"success": true, "path": "D:/DesktopPetData/Documents/xxx.docx", "title": "..."}
"""

import json
import sys
import os
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x86, 0xC1)
GRAY = RGBColor(0x66, 0x66, 0x66)
FONT_CN = "微软雅黑"


def sanitize_filename(name):
    return "".join(c for c in name if c not in '<>:"/\\|?*').strip() or "document"


def set_run_font(run, size, bold=False, color=None, italic=False):
    run.font.name = FONT_CN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # 中文字体（python-docx 需设置 eastAsia）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    sizes = {1: 20, 2: 16, 3: 13}
    size = sizes.get(level, 13)
    run = p.add_run(text)
    set_run_font(run, size, bold=True, color=PRIMARY if level <= 2 else ACCENT)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)  # 中文首行缩进2字符
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run, 11)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 11, italic=True, color=GRAY)
    p.paragraph_format.left_indent = Cm(1.0)
    return p


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"success": False, "error": "用法: python docx_gen.py <input.json> <output_dir>"}, ensure_ascii=False))
        return 1

    input_path = sys.argv[1]
    out_dir = sys.argv[2]

    try:
        # utf-8-sig 兼容 PowerShell 写入的 BOM，也兼容 Node 无 BOM 写入
        with open(input_path, "r", encoding="utf-8-sig") as f:
            desc = json.load(f)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"读取描述文件失败: {e}"}, ensure_ascii=False))
        return 1

    try:
        title = str(desc.get("title") or "文档")
        author = str(desc.get("author") or "")
        intro = str(desc.get("intro") or "")
        blocks = desc.get("blocks") or []

        doc = Document()
        # 默认样式
        style = doc.styles["Normal"]
        style.font.name = FONT_CN
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

        # 标题
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, 26, bold=True, color=PRIMARY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        if author or intro:
            meta = doc.add_paragraph()
            parts = []
            if author:
                parts.append(f"作者：{author}")
            if intro:
                parts.append(intro)
            meta_run = meta.add_run("    ".join(parts))
            set_run_font(meta_run, 10, color=GRAY)
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 分隔线（下边框）
        div = doc.add_paragraph()
        pPr = div._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pPr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "8",
            qn("w:space"): "1", qn("w:color"): "1F4E79"
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 内容块
        for b in blocks:
            t = str(b.get("type") or "p")
            text = str(b.get("text") or "")
            if t == "h1":
                add_heading(doc, text, 1)
            elif t == "h2":
                add_heading(doc, text, 2)
            elif t == "h3":
                add_heading(doc, text, 3)
            elif t == "bullet":
                add_bullet(doc, text)
            elif t == "number":
                add_number(doc, text)
            elif t == "quote":
                add_quote(doc, text)
            else:
                add_para(doc, text)

        os.makedirs(out_dir, exist_ok=True)
        safe_title = sanitize_filename(title)
        filename = f"{safe_title}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = os.path.join(out_dir, filename)
        doc.save(path)

        print(json.dumps({
            "success": True,
            "path": path.replace("\\", "/"),
            "title": title
        }, ensure_ascii=False))
        return 0
    except Exception as e:
        print(json.dumps({"success": False, "error": f"生成失败: {e}"}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    sys.exit(main())
